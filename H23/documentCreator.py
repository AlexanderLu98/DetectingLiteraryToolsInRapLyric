from docx import Document
from docx.shared import RGBColor
import colorsys
import re

# Your data
data = [['Living', 'off', 'borrowed', 'time', '0the0', '1clock1', '2tick2', '3faster3', "That'd", '4be4', '0the0', 'hour', '5they5', '1knock1', '0the0', '2slick2', '3blaster3', 'Dick', 'Dastardly', '6and6', 'Muttley', '7with7', '2sick2', '3laughter3', 'A', 'gunfight', '6and6', '5they5', '8come8', '9to9', '10cut10', '0the0', 'mixmaster', 'I-C-E', '11cold11', '12nice12', '9to9', '4be4', '11old11', 'Y2G', 'stee', '12twice12', '9to9', '11threefold11', 'He', '11sold11', '13scrolls13', 'lo', '6and6', '11behold11', 'Know', "who's", '0the0', 'illest', 'ever', '14like14', '0the0', '15greatest15', '16story16', '11told11', 'Keep', '17your17', '16glory16', '11gold11', '6and6', '3glitter3', 'For', '18half18', '18half18', '19of19', '20his20', "niggas'll", '21take21', 'him', '22out22', '0the0', '23picture23', 'The', '24other24', '18half18', '20is20', 'rich', '6and6', '25it25', "26don't26", '27mean27', 'shit-ta', 'Villain:', '0a0', '23mixture23', '27between27', 'both', '7with7', '0a0', '15twist15', '19of19', 'liquor', 'Chase', '25it25', '7with7', '17more17', 'beer', '15taste15', '25it25', '14like14', 'truth', '3or3', '29dare29', '30When30', '31he31', 'have', '0the0', '2mic2', "32it's32", '14like14', '0the0', 'place', '25get25', 'like:', 'Aw', 'yeah!', "33It's33", '14like14', '5they5', '34know34', "32what's32", "22'bout22", '9to9', '35happen35', 'Just', 'keep', 'ya', '36eye36', '22out22', '14like14', 'Aye', '36aye36', '35captain35', 'Is', '31he31', 'still', '0a0', '36fly36', '36guy36', '37clapping37', 'if', 'nobody', "ain't", '29hear29', '25it25', '38And38', '35can35', '5they5', 'testify', '8from8', 'in', '0the0', '25spirit25', 'Nah', 'In', '37living37', '0the0', 'true', '40gods40', 'Giving', "y'all", '37nothing37', '10but10', '0the0', '2lick2', '14like14', '9two9', 'broads', 'Got', '17more17', '41lyrics41', '35than35', '0the0', 'church', 'got', 'Ooh', 'Lord"s', '38And38', '31he31', '11hold11', '0the0', '2mic2', '6and6', '17your17', '35attention35', '14like14', '9two9', '40swords40', 'Either', '39that39', '3or3', '24either24', '43one43', '7with7', '9two9', '40blades40', '44on44', '25it25', 'Hey', '45you45', "26don't26", 'touch', '0the0', '2mic2', '14like14', "32it's32", 'AIDS', '44on44', '25it25', 'Yuck', "33It's33", '14like14', '0the0', '6end6', '9to9', '0the0', 'means', 'Fuck', 'type', '19of19', 'message', '39that39', '40sends40', '9to9', '0the0', '40fiends40', "That's", '36why36', '31he31', 'bring', '20his20', '26own26', '13needles13', '38And38', '25get25', '17more17', 'cheese', '35than35', 'Doritos', 'Cheetos', '3or3', 'Fritos', 'Slip', '14like14', 'Freudian', 'Your', '15first15', '6and6', 'last', 'step', '9to9', '37playing37', 'yourself', '14like14', '35accordion35', '30When30', '31he31', '42at42', '0the0', '2mic2', '45you45', "26don't26", 'go', '28next28', 'Leaving', 'pussy', '32cats32', '14like14', '36why36', 'hoes', '46need46', 'Kotex', 'Exercise', '41index41', "26won't26", '46need46', 'Bowflex', '38And38', "26won't26", '21take21', '0the0', '43one43', '7with7', '34no34', 'skinny', 'legs', '14like14', 'Joe', 'Tex']]

# Define the color mapping
color_mapping = {
    range(0, 18): [(n * 10) for n in range(0, 360) if n % 2 == 0],
    range(18, 37): [(n * 10) for n in range(0, 360) if n % 2 != 0]
}

# Create a new Word document
doc = Document()

# Iterate over the data
for line in data:
    # Create a new paragraph for each line
    paragraph = doc.add_paragraph()
    for word in line:
        # Check if the word has a rhyme number
        match = re.match(r'(\d+)(\w+)(\d+)', word)
        if match:
            # Split the word into the rhyme number and the actual word
            rhyme_number, actual_word, _ = match.groups()
            # Determine the color based on the rhyme number
            if int(rhyme_number) != -1:
                for key in color_mapping:
                    if int(rhyme_number) in key:
                        hue = color_mapping[key][int(rhyme_number) % len(color_mapping[key])]
                        # Convert the hue to RGB
                        r, g, b = [int(x * 255) for x in colorsys.hsv_to_rgb(hue / 360, 1, 1)]
                        color = RGBColor(r, g, b)
                        break
            run = paragraph.add_run(actual_word)
            run.font.color.rgb = color
        else:
            # If the word does not have a rhyme number, add it without any color
            run = paragraph.add_run(word)
        # Add a space after each word
        paragraph.add_run(' ')
    # Add a line break after each line
    doc.add_paragraph()

# Save the document
doc.save('MF DOOM - Accordion.docx')
