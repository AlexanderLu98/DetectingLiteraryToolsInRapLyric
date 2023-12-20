from docx import Document
from docx.shared import RGBColor
import colorsys
import re

# Insert data
data = [['I0', 'left', 'my1', 'legacy', 'hurt', "Fuckin'", 'absurd2', 'Like3', 'a4', 'shepherd', 'having5', 'sex6', 'with7', 'his8', 'sheep9', 'fuck', 'what10', 'you11', 'heard2', 'All', 'this12', 'talk', 'in13', 'my1', 'ear14', 'I0', 'got15', 'a4', 'idea4', 'Like3', 'the4', 'clerk', 'when13', 'you11', 'tryna', 'buy', 'beer14', 'ID', 'ya', 'Since', 'on17', 'the4', 'mic18', "I'm19", 'a4', 'nightmare', 'Fuck20', 'it21', 'I0', 'thought', 'this12', 'might22', 'be23', 'a4', 'good', 'time', 'to4', 'put', 'woke', 'me24', 'to4', 'rockabye', 'I0', 'got15', 'the4', 'bottle', 'of25', 'NyQuil', 'right22', 'here14', 'Right', 'here14', 'You', 'want26', 'the4', 'sleep9', 'me24', 'to4', 'wake27', 'you11', 'want26', 'Slim', 'Shady', 'EP', "That's", 'on17', 'the4', 'CD', 'cover', 'socking5', 'my1', 'mirror28', "Sockin'", 'my1', 'mirror28', 'I0', 'promise29', 'not15', 'to4', 'cry', 'crocodile', 'tears30', 'Crocodile', 'tears30', 'If', 'you11', 'end31', 'up32', 'shocked33', 'at34', 'my1', 'lyrics6', 'Shocked', 'at34', 'my1', 'lyrics6', 'Marshall', 'is8', 'dead35', 'in13', 'the4', 'water36', 'but10', 'not15', 'that34', 'I0', 'care37', 'Dre', 'said35', 'Rock', 'the4', 'boat', 'and31', 'the4', 'Doc', 'is8', 'my1', 'peer14', 'So', "it's38", 'unanimous29', "you're39", 'at34', 'attention', 'the4', "planet's38", 'listening5', 'And40', 'their37', 'banana4', 'splits38', 'again16', 'which', 'has41', 'its38', 'advantages8', 'But42', 'when13', 'you11', 'got15', 'nothing5', 'to4', 'say4', 'except43', 'for44', 'the4', 'hand31', 'your39', 'dick18', 'is8', 'in13', 'And40', 'if45', 'your39', "plan's46", 'to4', 'stick18', 'it21', 'in13', 'Janice', 'Dickinson', 'Imagine', 'if45', 'the4', 'Temazepam', 'is8', 'kicking5', 'in13', "it's38", 'having5', 'you11', 'panic-stricken', "You're", 'tripping5', 'off', 'of25', 'tryptophan', 'and31', 'tripped43', 'a4', 'fan', 'in13', 'Switzerland', 'Just', 'for44', 'asking5', 'to4', 'autograph', 'a4', 'picture', 'Then', 'ripped43', 'it21', 'in13', 'half', 'and31', 'whipped43', 'it21', 'at34', 'him', 'And40', 'kicked33', 'his8', 'ass', 'all', 'the4', 'way4', 'back47', 'to4', 'Michigan', 'But42', 'no48', 'matter36', 'how49', 'many', 'rounds46', 'or14', 'if45', 'I0', 'get21', 'knocked33', 'down50', 'In51', 'a4', 'bout52', 'and31', 'fell', 'to4', 'the4', 'ground31', 'I0', 'got15', 'a4', "fighter's53", 'mentality', "I'll54", 'get21', 'back47', 'up32', 'and31', 'fight22', 'with7', 'it21', 'In51', 'fact33', "I'ma", 'attacking5', 'the4', 'mic18', 'with7', 'it21', "I'll54", 'make27', 'it21', 'sound31', 'Sound', 'sound31', 'like', 'a4', "vampire's30", 'biting5', 'it21', "Bitin'", 'it21', 'But42', "I'd", 'have', 'to4', 'be23', "Dracula's", 'sidekick', 'Sidekick', 'To', 'be23', 'down50', 'for44', 'the4', 'count26', 'Haha', 'yeah', 'Yo', "album's", 'What55', 'out52', 'Yeah56', 'now', 'Uh57', 'Pow58', 'Pow58', 'wow', 'Woo', 'I0', "don't", 'Huh59', 'see', 'no48', 'Uh57', 'clouds', 'Nah', 'But42', 'aingt60', "gon'", 'No61', 'be23', 'no48', 'Huh59', 'drought52', 'Uh-uh', 'Smile', 'No61', 'frown50', 'Yeah56', 'upside', 'Ah', 'down50', 'Ya', 'Shut', 'my1', 'Nope', 'mouth', 'Yeah56', 'how49', 'How', 'That', 'aingt60', 'what10', 'What55', "I'm19", "'bout52", 'Shout62', 'Shout62', 'showers53', 'Fuck20', 'it21', 'ow48', 'ow48', "Thuggin'", 'thug', 'it21', 'out52', 'Ow', "'cause41"]]

# Ensure the color coding is easily to distingush from each other
color_mapping = {
    range(0, 18): [(n * 10) for n in range(0, 360) if n % 2 == 0],
    range(18, 37): [(n * 10) for n in range(0, 360) if n % 2 != 0]
}

doc = Document()

# Iterate over the data
for line in data:
    paragraph = doc.add_paragraph()
    for word in line:
        match = re.match(r'(\d+)(\w+)(\d+)', word)
        if match:
            rhyme_number, actual_word, _ = match.groups()
            # Determine the color based on the rhyme number
            if int(rhyme_number) != -1:
                for key in color_mapping:
                    if int(rhyme_number) in key:
                        hue = color_mapping[key][int(rhyme_number) % len(color_mapping[key])]
                        # Convert the hue to RGB
                        r, g, b = [int(x * 255) for x in colorsys.hsv_to_rgb(hue / 360, 1, 1)]
                        color = RGBColor(r, g, b)
                        break
            # Add the word to the paragraph with the determined color
            run = paragraph.add_run(actual_word)
            run.font.color.rgb = color
        else:
            # If the word does not have a rhyme number, add it without any color
            run = paragraph.add_run(word)
        paragraph.add_run(' ')
    doc.add_paragraph()
doc.save('Eminem - Rainy Days.docx')
